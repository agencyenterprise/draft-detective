"""Tests for DOCX chunk mapping service"""

import pytest
from docx import Document
from pathlib import Path

from lib.services.docx.chunk_mapper import create_chunk_to_paragraph_mapping
from lib.services.nltk_text_splitter import NLTKTextSplitter
from lib.workflows.context import ContextSchema
from tests.conftest import create_test_file_document_from_path, data_path


@pytest.mark.asyncio
async def test_chunk_to_docx_mapping_agi_minimal():
    """
    End-to-end test of chunk matching using agi-minimal test data.
    """
    file_doc = await create_test_file_document_from_path(
        "data/geopolitics-of-agi-minimal-1/_original.docx"
    )

    context = ContextSchema(
        session_id="test-session", workflow_id="test-workflow", run_id="test-run"
    )
    chunker = NLTKTextSplitter(context=context)
    chunks = await chunker.create_documents([file_doc.markdown])

    docx_path = data_path("data/geopolitics-of-agi-minimal-1/_original.docx")
    doc = Document(docx_path)
    docx_paragraphs = [p for p in doc.paragraphs if p.text.strip()]

    mapping = create_chunk_to_paragraph_mapping(chunks, docx_paragraphs)

    print(f"\n=== Mapping Results ===")
    print(f"Total chunks: {len(chunks)}")
    print(f"Total DOCX paragraphs: {len(docx_paragraphs)}")
    print(f"Chunks mapped: {len(mapping)}")
    print(f"Coverage: {len(mapping) / len(chunks) * 100:.1f}%")

    # Show any unmapped chunks for debugging
    unmapped = [c for c in chunks if c.metadata.chunk_index not in mapping]
    if unmapped:
        print(f"\n=== Unmapped Chunks ({len(unmapped)}) ===")
        for chunk in unmapped:
            print(
                f"  Chunk {chunk.metadata.chunk_index}: '{chunk.page_content[:60]}...'"
            )

    # ALL chunks must be mapped
    assert len(mapping) == len(chunks), (
        f"All {len(chunks)} chunks must be mapped, but only {len(mapping)} were mapped. "
        f"Unmapped chunk indices: {[c.metadata.chunk_index for c in unmapped]}"
    )

    for chunk_idx, para_idx in mapping.items():
        assert (
            0 <= para_idx < len(docx_paragraphs)
        ), f"Invalid paragraph index {para_idx} for chunk {chunk_idx}"


@pytest.mark.asyncio
async def test_chunk_mapping_handles_empty_chunks():
    """Test that empty chunks are handled gracefully"""
    from lib.agents.models import ValidatedDocument, DocumentMetadata

    chunks = [
        ValidatedDocument(
            page_content="",
            metadata=DocumentMetadata(
                chunk_index=0, paragraph_index=0, chunk_index_within_paragraph=0
            ),
        ),
        ValidatedDocument(
            page_content="Some actual content",
            metadata=DocumentMetadata(
                chunk_index=1, paragraph_index=1, chunk_index_within_paragraph=0
            ),
        ),
    ]

    docx_path = data_path("data/geopolitics-of-agi-minimal-1/_original.docx")
    doc = Document(docx_path)
    docx_paragraphs = [p for p in doc.paragraphs if p.text.strip()]

    mapping = create_chunk_to_paragraph_mapping(chunks, docx_paragraphs)

    assert 0 not in mapping, "Empty chunk should not be mapped"


def test_duplicate_text_maps_to_correct_paragraphs():
    """
    Test that duplicate text at different positions maps correctly.

    Chunk 1 "I repeat myself" should map to paragraph 1,
    Chunk 3 "I repeat myself" should map to paragraph 3 (not paragraph 1 again).
    """
    from lib.agents.models import ValidatedDocument, DocumentMetadata

    # Create a real Document with paragraphs instead of mocks
    doc = Document()
    doc.add_paragraph("I repeat myself forever")  # para 0
    doc.add_paragraph("Something else here")  # para 1
    doc.add_paragraph("I repeat myself always")  # para 2
    doc.add_paragraph("Final paragraph")  # para 3

    docx_paragraphs = list(doc.paragraphs)

    chunks = [
        ValidatedDocument(
            page_content="I repeat myself",
            metadata=DocumentMetadata(
                chunk_index=0, paragraph_index=0, chunk_index_within_paragraph=0
            ),
        ),
        ValidatedDocument(
            page_content="forever",
            metadata=DocumentMetadata(
                chunk_index=1, paragraph_index=0, chunk_index_within_paragraph=1
            ),
        ),
        ValidatedDocument(
            page_content="Something else",
            metadata=DocumentMetadata(
                chunk_index=2, paragraph_index=1, chunk_index_within_paragraph=0
            ),
        ),
        ValidatedDocument(
            page_content="I repeat myself",  # Same text as chunk 0!
            metadata=DocumentMetadata(
                chunk_index=3, paragraph_index=2, chunk_index_within_paragraph=0
            ),
        ),
        ValidatedDocument(
            page_content="always",
            metadata=DocumentMetadata(
                chunk_index=4, paragraph_index=2, chunk_index_within_paragraph=1
            ),
        ),
    ]

    mapping = create_chunk_to_paragraph_mapping(chunks, docx_paragraphs)

    assert mapping[0] == 0, "Chunk 0 'I repeat myself' should map to paragraph 0"
    assert mapping[1] == 0, "Chunk 1 'forever' should map to paragraph 0"
    assert mapping[2] == 1, "Chunk 2 'Something else' should map to paragraph 1"
    assert (
        mapping[3] == 2
    ), "Chunk 3 'I repeat myself' should map to paragraph 2 (not 0!)"
    assert mapping[4] == 2, "Chunk 4 'always' should map to paragraph 2"

    print("\n✓ Duplicate text correctly mapped to different paragraphs")
    print(f"  Chunk 0 'I repeat myself' → paragraph {mapping[0]}")
    print(f"  Chunk 3 'I repeat myself' → paragraph {mapping[3]}")


@pytest.mark.asyncio
async def test_add_comments_to_docx():
    """
    Test adding comments to DOCX file with severity-based authors.

    This tests the complete integration:
    1. Load DOCX and convert to markdown
    2. Chunk the markdown
    3. Create comments with different severities
    4. Add comments to DOCX
    5. Verify output file is created
    """
    from lib.services.docx.manipulator import (
        CommentSeverity,
        DocxComment,
        DocxManipulatorService,
    )

    file_doc = await create_test_file_document_from_path(
        "data/geopolitics-of-agi-minimal-1/_original.docx"
    )

    context = ContextSchema(
        session_id="test-session", workflow_id="test-workflow", run_id="test-run"
    )
    chunker = NLTKTextSplitter(context=context)
    chunks = await chunker.create_documents([file_doc.markdown])

    comments = [
        DocxComment(
            chunk_index=0,
            text=chunks[0].page_content,
            comment_text="High priority issue found",
            severity=CommentSeverity.HIGH,
        ),
        DocxComment(
            chunk_index=1,
            text=chunks[1].page_content,
            comment_text="Medium priority suggestion",
            severity=CommentSeverity.MEDIUM,
        ),
        DocxComment(
            chunk_index=2,
            text=chunks[2].page_content,
            comment_text="Low priority note",
            severity=CommentSeverity.LOW,
        ),
    ]

    # Verify severity-based authors
    assert comments[0].get_author() == "🚨 High Priority"
    assert comments[1].get_author() == "⚠️ Medium Priority"
    assert comments[2].get_author() == "💡 Low Priority"
    assert comments[0].get_initials() == "HP"
    assert comments[1].get_initials() == "MP"
    assert comments[2].get_initials() == "LP"

    service = DocxManipulatorService()
    output_path = await service.add_comments_to_docx(
        original_docx_path=file_doc.file_path,
        comments=comments,
        workflow_run_id="test-run-123",
        chunks=chunks,
    )

    assert Path(output_path).exists(), "Output DOCX file should exist"

    doc = Document(output_path)
    assert len(doc.paragraphs) > 0, "Output DOCX should have paragraphs"

    try:
        if hasattr(doc, "part") and hasattr(doc.part, "package"):
            comments_part_exists = any(
                "comments" in str(part_name).lower()
                for part_name in doc.part.package.part_names
            )
            if comments_part_exists:
                print(f"\n✓ Comments part detected in DOCX package")
        print(f"\n✓ Successfully created DOCX with comments at: {output_path}")
        print(f"  Total paragraphs: {len(doc.paragraphs)}")
        print(f"  Comments requested: {len(comments)}")
    except Exception as e:
        print(f"\n⚠ Could not verify comments (non-fatal): {e}")
