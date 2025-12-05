import logging
from enum import StrEnum
from pathlib import Path
from typing import List, Optional

from docx import Document
from pydantic import BaseModel

from lib.agents.models import ValidatedDocument
from lib.config.env import config
from lib.services.docx_chunk_mapper import create_chunk_to_paragraph_mapping

logger = logging.getLogger(__name__)


class CommentSeverity(StrEnum):
    """Severity levels for DOCX comments."""

    NONE = "none"
    LOW = "low"
    MEDIUM = "medium"
    HIGH = "high"


# Severity to author name and icon mapping
SEVERITY_AUTHORS = {
    CommentSeverity.HIGH: ("🚨 High Priority", "HP"),
    CommentSeverity.MEDIUM: ("⚠️ Medium Priority", "MP"),
    CommentSeverity.LOW: ("💡 Low Priority", "LP"),
    CommentSeverity.NONE: ("📝 Note", "NT"),
}


class DocxComment(BaseModel):
    """Represents a comment to be added to a docx file"""

    chunk_index: int
    text: str
    comment_text: str
    severity: Optional[CommentSeverity] = None
    author: Optional[str] = None  # If not provided, derived from severity

    def get_author(self) -> str:
        """Get author name, derived from severity if not explicitly set."""
        if self.author:
            return self.author
        if self.severity:
            return SEVERITY_AUTHORS.get(
                self.severity, SEVERITY_AUTHORS[CommentSeverity.NONE]
            )[0]
        return "AI Reviewer"

    def get_initials(self) -> str:
        """Get initials for the comment author."""
        if self.severity and not self.author:
            return SEVERITY_AUTHORS.get(
                self.severity, SEVERITY_AUTHORS[CommentSeverity.NONE]
            )[1]
        # Derive from author name
        author = self.get_author()
        # Skip emoji at start if present
        parts = author.split()
        # Filter out emoji-only parts
        text_parts = [
            p for p in parts if p.isalpha() or (len(p) > 1 and p[-1].isalpha())
        ]
        if len(text_parts) >= 2:
            return f"{text_parts[0][0]}{text_parts[1][0]}".upper()
        elif len(text_parts) == 1:
            return text_parts[0][:2].upper()
        return "AI"


class DocxManipulatorService:
    """Service for manipulating DOCX files with AI-generated comments"""

    SUPPORTED_EXTENSIONS = {".docx"}

    def get_output_path(self, workflow_run_id: str) -> Path:
        """Get the deterministic output path for a processed docx file"""
        output_dir = Path(config.FILE_UPLOADS_MOUNT_PATH) / "processed_docx"
        output_dir.mkdir(exist_ok=True)
        return output_dir / f"{workflow_run_id}_reviewed.docx"

    async def add_comments_to_docx(
        self,
        original_docx_path: str,
        comments: List[DocxComment],
        workflow_run_id: str,
        chunks: List[ValidatedDocument] | None = None,
    ) -> str:
        """
        Add comments to a DOCX file based on chunk indices.

        Args:
            original_docx_path: Path to the original .docx file
            comments: List of comments to add
            workflow_run_id: Workflow run ID for deterministic naming
            chunks: Optional list of document chunks for mapping

        Returns:
            Path to the updated .docx file
        """
        original_path = Path(original_docx_path)

        if original_path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"File must be .docx, got {original_path.suffix}")

        if not original_path.exists():
            raise FileNotFoundError(f"Original file not found: {original_docx_path}")

        output_path = self.get_output_path(workflow_run_id)

        logger.info(
            f"Creating reviewed docx at {output_path} with {len(comments)} comments"
        )

        doc = Document(original_docx_path)
        docx_paragraphs = [p for p in doc.paragraphs if p.text.strip()]

        chunk_to_para_mapping = {}
        if chunks:
            chunk_to_para_mapping = create_chunk_to_paragraph_mapping(
                chunks, docx_paragraphs
            )
        else:
            logger.warning("No chunks provided, comments will not be added")

        comments_added = 0
        comments_skipped = 0

        for comment in comments:
            para_idx = chunk_to_para_mapping.get(comment.chunk_index)

            if para_idx is None:
                logger.warning(
                    f"No paragraph mapping found for chunk {comment.chunk_index}, "
                    f"comment '{comment.comment_text[:50]}...' will be skipped"
                )
                comments_skipped += 1
                continue

            if para_idx >= len(docx_paragraphs):
                logger.warning(
                    f"Invalid paragraph index {para_idx} for chunk {comment.chunk_index}"
                )
                comments_skipped += 1
                continue

            try:
                paragraph = docx_paragraphs[para_idx]
                self._add_comment_to_paragraph(
                    doc,
                    paragraph,
                    comment.comment_text,
                    comment.get_author(),
                    comment.get_initials(),
                )
                comments_added += 1
            except Exception as e:
                logger.error(
                    f"Failed to add comment to paragraph {para_idx}: {e}", exc_info=True
                )
                comments_skipped += 1

        doc.save(str(output_path))

        logger.info(
            f"Successfully created reviewed docx: {comments_added} comments added, "
            f"{comments_skipped} skipped"
        )
        return str(output_path)

    def _add_comment_to_paragraph(
        self,
        doc: Document,
        paragraph,
        comment_text: str,
        author: str,
        initials: str,
    ):
        """
        Add a comment to a paragraph.

        Args:
            doc: The Document object
            paragraph: The paragraph to add comment to
            comment_text: The comment text
            author: The comment author
            initials: The author initials
        """
        try:
            doc.add_comment(
                runs=paragraph.runs if paragraph.runs else [paragraph.add_run("")],
                text=comment_text,
                author=author,
                initials=initials,
            )
        except AttributeError as e:
            logger.warning(f"Error inserting comment to docx: {e}")
            raise


docx_manipulator_service = DocxManipulatorService()
